from random import randint
import random
import string
import copy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx import Document
from docx.shared import Cm, Inches, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.dml.color import ColorFormat
import hgtk
import requests
from bs4 import BeautifulSoup
import re
import os
from pathlib import Path


class MakeWordSearch():
    def __init__(self, word_image, width, height, diff, option, pic_on, korean, chosung):
        try:
            self.desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'desktop')
            os.listdir(self.desktop)
        except FileNotFoundError:
            self.desktop = '.'
        self.word_image = word_image
        self.width = width
        self.height = height
        self.diff = diff
        self.option = option
        self.pic_on = pic_on
        self.korean = korean
        self.chosung = chosung

    def show_me_the_way(self, right, up):
        start_x_code = ''
        start_y_code = ''
        direction_x_code = ''
        direction_y_code = ''
        if right == 1:
            start_x_code = "randint(1, self.width-len(word)+1)-1"
            direction_x_code = "startpoint[0] + i"
        elif right == 0:
            start_x_code = "randint(1, self.width)-1"
            direction_x_code = "startpoint[0]"
        elif right == -1:
            start_x_code = "randint(len(word)+1, self.width)-1"
            direction_x_code = "startpoint[0] - i"
        if up == 1:
            start_y_code = "randint(len(word)+1, self.height)-1"
            direction_y_code = "startpoint[1] - i"
        elif up == 0:
            start_y_code = "randint(1, self.height)-1"
            direction_y_code = "startpoint[1]"
        elif up == -1:
            start_y_code = "randint(1, self.height-len(word)+1)-1"
            direction_y_code = "startpoint[1] + i"
        return [start_x_code, start_y_code, direction_x_code, direction_y_code]

    def word_exist(self, word, place_num):
        word_num = len(word)
        word_zero = ""
        for i in range(word_num):
            word_zero = word_zero + '0'
        word_in_puzzle = ""
        for coordinate in place_num:
            x = coordinate[0]
            y = coordinate[1]
            word_in_puzzle += str(self.puzzle_origin[y][x])
        if word_zero == word_in_puzzle:
            return True
        else:

            for i in range(word_num):
                if word_in_puzzle[i] != word[i] and word_in_puzzle[i] != '0':
                    return False
            return True

    #가로 순방향으로 만듦
    def col(self, word):
        right = 1
        up =0
        word_set = self.show_me_the_way(right, up) #0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]),eval(word_set[3])])


        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)
    #가로 역방향으로 만듦
    def col_rev(self, word):
        right = -1
        up = 0
        word_set = self.show_me_the_way(right, up)  # 0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]), eval(word_set[3])])


        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)
    #세로 순방향으로 만듦
    def row(self, word):
        right = 0
        up = 1
        word_set = self.show_me_the_way(right, up)  # 0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]), eval(word_set[3])])


        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)
    #세로 역방향으로 만듦
    def row_rev(self, word):
        right = 0
        up = -1
        word_set = self.show_me_the_way(right, up)  # 0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]), eval(word_set[3])])

        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)

    #대각선 위 순방향
    def diagup(self, word):
        right = 1
        up = 1
        word_set = self.show_me_the_way(right, up)  # 0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]), eval(word_set[3])])


        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)
    #대각선 아래 순항뱡
    def diagdown(self, word):
        right = 1
        up = -1
        word_set = self.show_me_the_way(right, up)  # 0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]), eval(word_set[3])])


        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)
    #대각선 위 역방향
    def diagup_rev(self, word):
        right = -1
        up = 1
        word_set = self.show_me_the_way(right, up)  # 0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]), eval(word_set[3])])


        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)
    #대각선 아래 역방향
    def diagdown_rev(self, word):
        right = -1
        up = -1
        word_set = self.show_me_the_way(right, up)  # 0:start_x_code, 1:start_y_code, 2:direction_x_code, 3:direction_y_code
        startpoint = [eval(word_set[0]), eval(word_set[1])]

        place_num = []
        for i in range(len(word)):
            place_num.append([eval(word_set[2]), eval(word_set[3])])


        if self.word_exist(word, place_num):
            for i, letter in enumerate(word):
                self.puzzle_origin[eval(word_set[3])][eval(word_set[2])] = letter
        else:
            exec(self.difficulty)

        return


    def make_puzzle(self):
        difficulty = self.diff
        option = self.option
        if difficulty == 1:
            self.difficulty = 'random.choice([self.col, self.row_rev])(word)'
        elif difficulty == 2:
            self.difficulty = "random.choice([self.col, self.col_rev, self.row, self.row_rev])(word)"
        elif difficulty == 3:
            self.difficulty = "random.choice([self.col, self.row_rev, self.diagup, self.diagdown])(word)"
        elif difficulty == 4:
            self.difficulty = "random.choice([self.col, self.col_rev, self.row, self.row_rev, self.diagup, self.diagdown])(word)"
        elif difficulty == 5:
            self.difficulty = "random.choice([self.col, self.col_rev, self.row, self.row_rev, self.diagup, self.diagup_rev, self.diagdown, self.diagdown_rev])(word)"
        elif difficulty == 6:
            self.difficulty = "random.choice([self.col, self.col_rev, self.row, self.row_rev, self.diagup, self.diagup_rev, self.diagdown, self.diagdown_rev])(word)"
        self.puzzle_origin = []
        for i in range(self.height):
            self.puzzle_origin.append([])
            for j in range(self.width):
                self.puzzle_origin[i].append('0')
        print("퍼즐 만드는 중")
        words = [word[0] for word in self.word_image]
        for word in words:
            exec(self.difficulty)

        string_words = ''.join(words)
        from collections import Counter
        count_alpha = Counter(string_words)

        common_alph = ''
        for alph in count_alpha.most_common(5):
            common_alph += alph[0]

        data = ''
        if self.korean:
            f = open("random_words.txt", 'r')
            data = f.read()
            regex_f = r'[가-힣]+'
            search_target_f = data
            data = ''.join(list(set(re.findall(regex_f, search_target_f))))

        printed_words = ''
        puzzle = copy.deepcopy(self.puzzle_origin)
        for i in range(self.height):
            for j in range(self.width):
                if self.puzzle_origin[i][j] == "0":
                    fill_alph = random.choice(string.ascii_lowercase)
                    if self.korean:
                        fill_alph = random.choice(data)
                    #글자들 되도록 겹치지 않게 하기 위해서 많이 나오는 글자 한번쯤은 피할 수 있도록 한다.
                    if option == 0:
                        puzzle[i][j] = fill_alph
                    elif option == 1:
                        if fill_alph in common_alph:
                            fill_alph = random.choice(string.ascii_lowercase)
                            if self.korean:
                                fill_alph = random.choice(data)
                        puzzle[i][j] = fill_alph
                        printed_words += puzzle[i][j]
                    #글자가 겹치도록 하기 위해서 많이 나온 글자와 무작위 글자들 중에서 고르도록 한다.
                    elif option == 2:
                        common_alph_list = []
                        puzzle[i][j] = random.choice([fill_alph, random.choice(count_alpha.most_common(7))[0]])
                    printed_words += puzzle[i][j]



        # write to docx file
        # Write to docx to puzzle.docx
        document = Document()
        #changing the page margins
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(0.8)
            section.left_margin = Cm(2.3)
            section.right_margin = Cm(2.3)
        heading = 'Word Puzzle'
        if self.korean:
            heading = "낱말 찾기"
        document.add_heading(heading, 0)

        para_belong = document.add_paragraph(' 학년  반')
        para_belong.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        para_name = document.add_paragraph('이름: ')
        para_name.add_run("홍길동").font.color.rgb = RGBColor(255, 255, 255)
        para_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        puzzle_table = document.add_table(rows=self.height, cols=self.width, style='Table Grid')
        puzzle_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        self.set_height = 7500 / self.height
        for i, row in enumerate(puzzle_table.rows):
            #######################세로 길이 정하기!
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(self.set_height))
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                #####가로 길이 정하기!
                cell.width = Inches(5)
                cell.text = puzzle[i][j]
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.style.font.bold = True
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcPr.append(tcVAlign)

        if self.pic_on:
            word_num = len(words)
            if word_num <= 15:
                size = 5
            elif word_num <= 21:
                size = (word_num+2)//3
            else:
                size = 7
            hint_table = document.add_table(rows = (len(words)+size-1)//size * 2, cols = size, style = 'Table Grid')
            hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, row in enumerate(hint_table.rows):
                #######################세로 길이 정하기!
                if i%2 == 0:
                    # accessing row xml and setting tr height
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), '1000')
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)
                elif i%2 == 1:
                    # accessing row xml and setting tr height
                    tr = row._tr
                    trPr = tr.get_or_add_trPr()
                    trHeight = OxmlElement('w:trHeight')
                    trHeight.set(qn('w:val'), '60')
                    trHeight.set(qn('w:hRule'), "atLeast")
                    trPr.append(trHeight)

                for j, cell in enumerate(row.cells):
                    index = i//2*size + j
                    #####가로 길이 정하기!
                    cell.width = Inches(10)

                    #단어 수 만큼 반복하기
                    if index < len(words):
                        if i % 2 == 1:
                            cell.text = words[index]
                            if self.chosung:
                                word = words[index]
                                cho_word = ''
                                for chr in word:
                                    chosung = hgtk.letter.decompose(chr)[0]
                                    cho_word += chosung
                                cell.text = cho_word
                        for paragraph in cell.paragraphs:
                            if i % 2 == 0:
                                if self.word_image[index][1] == "None":
                                    cell.text = "사진 없음"
                                else:
                                    try:
                                        run = paragraph.add_run()
                                        if size > 8:
                                            run.add_picture(self.word_image[index][1], width=Mm(15), height=Mm(15))
                                        else:
                                            run.add_picture(self.word_image[index][1], width=Mm(40 - size * 3.2),
                                                            height=Mm(40 - size * 3.2))
                                    except:
                                        pass


                            #####가운데 정렬!!
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            paragraph.style.font.bold = True
                    #####상하 방향에서 가운데 정렬
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), "center")
                    tcPr.append(tcVAlign)
        else:
            if not self.korean:
                hint_table = document.add_table(rows=1, cols=1, style='Table Grid')
                hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hint_table_row = hint_table.rows[0]
                hint_tr = hint_table_row._tr
                hint_trPr = hint_tr.get_or_add_trPr()
                hint_trHeight = OxmlElement('w:trHeight')
                hint_trHeight.set(qn('w:val'), '1000')
                hint_trHeight.set(qn('w:hRule'), "atLeast")
                hint_trPr.append(hint_trHeight)
                hint_table_cell = hint_table_row.cells[0]
                hint = ''
                parenthesis = re.compile(r'(\s)?\(.*\)(\s)?')
                bracket = re.compile(r'(\s)?\[.*\](\s)?')
                for word in words:
                    print("사전에 찾는중... " + word)
                    req = requests.get('http://endic.naver.com/small_search.nhn?query=' + word)
                    html = req.text
                    soup = BeautifulSoup(html, 'html.parser')
                    meanings = soup.select('span.fnt_k05')
                    if meanings:
                        text = meanings[0].text
                        text = re.sub(parenthesis, '', text)
                        text = re.sub(bracket, '', text)
                        print(text)
                        hint += word + "({})".format(text) + ', '
                hint_table_cell.width = Inches(100)
                for paragraph in hint_table_cell.paragraphs:
                    paragraph.add_run(hint.strip(', '))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                tc = hint_table_cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcPr.append(tcVAlign)

            else:
                if self.chosung:
                    hint_table = document.add_table(rows=1, cols=1, style='Table Grid')
                    hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hint_table_row = hint_table.rows[0]
                    hint_tr = hint_table_row._tr
                    hint_trPr = hint_tr.get_or_add_trPr()
                    hint_trHeight = OxmlElement('w:trHeight')
                    hint_trHeight.set(qn('w:val'), '1000')
                    hint_trHeight.set(qn('w:hRule'), "atLeast")
                    hint_trPr.append(hint_trHeight)
                    hint_table_cell = hint_table_row.cells[0]
                    hint = ''
                    for word in words:
                        cho_word = ''
                        for chr in word:
                            chosung = hgtk.letter.decompose(chr)[0]
                            cho_word += chosung
                        hint += cho_word + ', '
                    hint_table_cell.width = Inches(100)
                    for paragraph in hint_table_cell.paragraphs:
                        paragraph.add_run(hint.strip(', '))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    tc = hint_table_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), "center")
                    tcPr.append(tcVAlign)
                else:
                    hint_table = document.add_table(rows=1, cols=1, style='Table Grid')
                    hint_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    hint_table_row = hint_table.rows[0]
                    hint_tr = hint_table_row._tr
                    hint_trPr = hint_tr.get_or_add_trPr()
                    hint_trHeight = OxmlElement('w:trHeight')
                    hint_trHeight.set(qn('w:val'), '1000')
                    hint_trHeight.set(qn('w:hRule'), "atLeast")
                    hint_trPr.append(hint_trHeight)
                    hint_table_cell = hint_table_row.cells[0]
                    hint = ''
                    for word in words:
                        hint += word + ', '
                    hint_table_cell.width = Inches(100)
                    for paragraph in hint_table_cell.paragraphs:
                        paragraph.add_run(hint.strip(', '))
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    tc = hint_table_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), "center")
                    tcPr.append(tcVAlign)





        # 정답 파일 쓰기
        answ_doc = Document()
        answer_table = answ_doc.add_table(rows=self.height, cols=self.width, style='Table Grid')
        answer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, row in enumerate(answer_table.rows):
            #######################세로 길이 정하기!
            # accessing row xml and setting tr height
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), str(self.set_height))
            trHeight.set(qn('w:hRule'), "atLeast")
            trPr.append(trHeight)

            for j, cell in enumerate(row.cells):
                #####가로 길이 정하기!
                cell.width = Inches(8)
                cell.text = self.puzzle_origin[i][j]
                for paragraph in cell.paragraphs:
                    #####가운데 정렬!!
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.style.font.bold = True
                    if cell.text == '0':
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                    else:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(255, 0, 0)
                #####상하 방향에서 가운데 정렬
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcVAlign = OxmlElement('w:vAlign')
                tcVAlign.set(qn('w:val'), "center")
                tcPr.append(tcVAlign)

        answ_doc.save(str(self.desktop) + '\puzzle_정답.hwp')
        document.save(str(self.desktop) +'\puzzle.docx')
        document.save(str(self.desktop) +'\puzzle.hwp')
        print("바탕화면에 puzzle.docx와 puzzle.hwp 로 저장")